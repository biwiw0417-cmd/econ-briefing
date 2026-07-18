# -*- coding: utf-8 -*-
"""브리핑 JSON → Word(.docx) 파일 생성.

data/briefings/*.json 중 docs/files/에 docx가 없는 날짜를 모두 생성한다.
서식: 제목(날짜+주제) → 뉴스 → 경제 지식 → 용어 → 복습 → 퀴즈(정답은 마지막 페이지)
"""
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
BRIEFINGS = ROOT / "data" / "briefings"
OUT_DIR = ROOT / "docs" / "files"

ACCENT = RGBColor(0x2F, 0x6F, 0x5E)  # 은은한 초록
GRAY = RGBColor(0x88, 0x88, 0x88)


def style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5


def add_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(18)


def add_body(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
    p.add_run(text)


def build_one(briefing: dict, out_path: Path) -> None:
    doc = Document()
    style_base(doc)

    date = briefing["date"]
    lesson = briefing["lesson"]

    # 제목
    title = doc.add_heading(f"{date} 아침 경제 브리핑", level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(f"오늘의 주제: {lesson['topic']}")
    run.font.color.rgb = GRAY

    # ① 뉴스
    add_label(doc, "① 오늘의 경제 뉴스")
    for i, news in enumerate(briefing["news"], 1):
        h = doc.add_paragraph()
        r = h.add_run(f"{i}. {news['title']}")
        r.bold = True
        r.font.size = Pt(12)
        if news.get("source_name"):
            h.add_run(f"  ({news['source_name']})").font.color.rgb = GRAY
        add_body(doc, news["summary"])
        add_body(doc, news["why_it_matters"], "왜 중요한가요?  ")
        if news.get("link_to_lesson"):
            add_body(doc, news["link_to_lesson"], "오늘 배움과의 연결  ")

    # ② 경제 지식
    unit_label = "주간 복습" if lesson["unit"] == 0 else f"Unit {lesson['unit']}"
    add_label(doc, f"② 오늘의 경제 지식 — [{unit_label}] {lesson['topic']}")
    add_body(doc, lesson["body"])
    if lesson.get("news_connection"):
        add_body(doc, lesson["news_connection"], "오늘 뉴스와의 연결  ")

    # ③ 용어
    term = briefing["term"]
    add_label(doc, f"③ 오늘의 용어 — {term['term']}")
    add_body(doc, term["definition"], "정의  ")
    add_body(doc, term["analogy"], "비유  ")
    add_body(doc, term["usage"], "사용 예  ")

    # ④ 복습
    add_label(doc, "④ 복습 코너")
    if briefing["review"]:
        for rv in briefing["review"]:
            add_body(doc, rv["refresher"],
                     f"{rv['term']} (배운 날: {rv.get('learned_date', '')})  ")
    else:
        add_body(doc, "아직 복습할 용어가 없습니다. 용어가 쌓이면 3일/7일 간격으로 다시 만나요.")

    # ⑤ 퀴즈 (문제만)
    add_label(doc, "⑤ 확인 퀴즈")
    for i, q in enumerate(briefing["quiz"], 1):
        add_body(doc, q["question"], f"Q{i}. ")
        for j, choice in enumerate(q["choices"], 1):
            doc.add_paragraph(f"    {j}) {choice}")

    # 정답·해설은 페이지 넘겨서
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_label(doc, "퀴즈 정답 및 해설")
    for i, q in enumerate(briefing["quiz"], 1):
        add_body(doc, f"{q['answer_index'] + 1}) {q['choices'][q['answer_index']]}",
                 f"Q{i} 정답  ")
        add_body(doc, q["explanation"], "해설  ")

    doc.save(out_path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    built = 0
    for path in sorted(BRIEFINGS.glob("*.json")):
        briefing = json.loads(path.read_text(encoding="utf-8"))
        out_path = OUT_DIR / f"{briefing['date']}_경제브리핑.docx"
        if not out_path.exists():
            build_one(briefing, out_path)
            print(f"[docx] 생성: {out_path.name}")
            built += 1
    if built == 0:
        print("[docx] 새로 만들 파일 없음")


if __name__ == "__main__":
    main()
